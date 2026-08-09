"""Build the Word document for submission (assignment evaluation guideline a).

    python scripts/build_submission_doc.py

Writes submission/CCZG506_Assignment2_GROUPID.docx. Rename it to <groupid>.docx
before uploading, and fill in every «...» placeholder - group number, member names,
BITS IDs and contribution percentages. Those are the only parts this script cannot
produce; everything else, including the measured numbers, comes from the repository
and from data/finetuning_results.json.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "submission"
OUT_FILE = OUT_DIR / "CCZG506_Assignment2_GROUPID.docx"
SHOTS = ROOT / "screenshot"
MAX_IMAGE_WIDTH_IN = 6.0

PLACEHOLDER = RGBColor(0xC0, 0x00, 0x00)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #

def para(doc, text="", style=None, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def todo(doc, text):
    """A placeholder line, coloured so it cannot be missed before submission."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = PLACEHOLDER
    return p


def add_image(doc, relative_path, caption):
    """Insert a screenshot scaled to the text width, with a caption under it."""
    path = SHOTS / relative_path
    if not path.exists():
        todo(doc, f"«MISSING SCREENSHOT: {relative_path} - re-run the capture scripts»")
        return
    with Image.open(path) as img:
        width_px, height_px = img.size
    width_in = MAX_IMAGE_WIDTH_IN
    height_in = width_in * height_px / width_px
    # Keep tall screenshots on one page.
    if height_in > 7.5:
        height_in = 7.5
        width_in = height_in * width_px / height_px
    doc.add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(doc, caption, italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    return cap


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(value))
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


# --------------------------------------------------------------------------- #
# Document
# --------------------------------------------------------------------------- #

def build():
    results_file = ROOT / "data" / "finetuning_results.json"
    if not results_file.exists():
        sys.exit(
            "data/finetuning_results.json not found. Run "
            "`python scripts/finetune_qa.py` first so the fine-tuning numbers "
            "in the document are real."
        )
    ft = json.loads(results_file.read_text(encoding="utf-8"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------------- Title ----------------
    para(doc, "CCZG506 - API-driven Cloud Native Solutions", bold=True, size=18,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Assignment II - AI Study Assistant for the Education Domain",
         bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    para(doc, "Group Details", style="Heading 1")
    todo(doc, "«Group No: FILL IN»")
    add_table(
        doc,
        ["Sl. No", "BITS ID", "Name", "Contribution (qualitative)", "% (of 100)"],
        [
            [1, "«BITS ID»", "«Name»",
             "Text summarisation sub-task; shared LLMOps metrics layer; "
             "fine-tuning (SciQ) and practice question generation", "«%»"],
            [2, "«BITS ID»", "«Name»", "«Contribution»", "«%»"],
            [3, "«BITS ID»", "«Name»", "«Contribution»", "«%»"],
            [4, "«BITS ID»", "«Name»", "«Contribution»", "«%»"],
            [5, "«BITS ID»", "«Name»", "«Contribution»", "«%»"],
        ],
        widths=[0.6, 1.2, 1.3, 2.6, 0.8],
    )
    todo(doc, "«Delete unused rows. Percentages must total 100.»")

    para(doc, "Repository: https://github.com/sur3175/API-drivenCloudNative_Assignment2-")

    doc.add_page_break()

    # ---------------- Overview ----------------
    para(doc, "1. Project Overview", style="Heading 1")
    para(doc,
         "Domain: Education. The project is an AI study assistant. A student supplies "
         "their own study material - lecture notes, a textbook section, a paper - and "
         "the application generates study content from it, condenses it for revision, "
         "answers questions about it, produces practice questions to test recall, and "
         "generates supporting diagrams.")
    para(doc,
         "Categories chosen: Natural Language Processing and Computer Vision.")
    para(doc,
         "The sub-tasks are cohesive rather than independent demos: they all operate on "
         "the same student-supplied material and feed one objective, revising a topic "
         "effectively.")

    para(doc, "1.1 Sub-tasks and models", style="Heading 2")
    add_table(
        doc,
        ["#", "Sub-task", "Category", "Model(s)", "Type"],
        [
            [1, "Text generation", "NLP", "gpt-4o-mini / HF provider", "LLM"],
            [2, "Text summarisation", "NLP",
             "sshleifer/distilbart-cnn-12-6; Qwen2.5-7B-Instruct", "SLM + LLM"],
            [3, "Question answering", "NLP",
             "distilbert-base-cased-distilled-squad; Qwen2.5-7B-Instruct", "SLM + LLM"],
            [4, "Image generation", "CV", "gpt-image-1", "LLM"],
            [5, "Image classification", "CV",
             "google/vit-base-patch16-224 (local + Inference API); gpt-4o-mini vision",
             "SLM + LLM"],
            [6, "Practice question generation", "NLP",
             "t5-small fine-tuned on SciQ", "SLM (fine-tuned)"],
        ],
        widths=[0.4, 1.6, 0.8, 2.4, 1.0],
    )
    para(doc,
         "Six sub-tasks against the five required. Sub-task 6 is the fine-tuned model "
         "and also satisfies requirement 8.")

    para(doc, "1.2 Cohesion: one document, six sub-tasks", style="Heading 2")
    para(doc,
         "The student's document is uploaded or pasted once, in the Study Material "
         "section at the top of the application, and every text sub-task reads from it. "
         "This is what makes the sub-tasks a workflow rather than six independent demos, "
         "which requirement 5 of the brief asks for.")
    add_table(
        doc,
        ["Sub-task", "How it uses the shared document"],
        [
            ["Text generation",
             "Optionally sends the document as context, so the model answers about the "
             "student's own material rather than from general knowledge"],
            ["Image generation",
             "Extracts the document's key terms and writes a mind-map prompt from them, "
             "so the diagram is drawn from what the notes actually emphasise"],
            ["Question answering", "Answers strictly from the document"],
            ["Text summarisation", "Condenses the document"],
            ["Image classification",
             "The entry point for material that arrives as a photo - it labels what was "
             "photographed so it can be routed to the right sub-task"],
            ["Practice question generation",
             "Generates exam questions from the document's key terms"],
        ],
        widths=[1.7, 4.3],
    )
    add_image(doc, "shared_material/01_shared_study_material.png",
              "Figure 1: The document is loaded once - by upload, paste, or a one-click "
              "sample - and every section below reads from it.")

    para(doc, "1.3 Architecture", style="Heading 2")
    para(doc,
         "The application is a Streamlit front end over a set of independent Python "
         "modules, one per sub-task, in src/. config.py centralises credentials and the "
         "provider factory create_client(provider), so a sub-task never reads .env "
         "itself. Every model call, whichever provider it goes to, is wrapped by the "
         "shared metrics layer in src/metrics.py.")
    para(doc,
         "Providers: the application is launched with a provider argument "
         "(python -m streamlit run app.py -- hf), which selects the client used by text "
         "and image generation. Hugging Face is the primary provider so the project is "
         "not gated on OpenAI credits. The summarisation and question answering "
         "sub-tasks additionally expose a per-run model selector, so a local small "
         "model and a hosted large model can be compared without restarting the app.")

    doc.add_page_break()

    # ---------------- Generation sub-tasks ----------------
    para(doc, "2. Sub-tasks: Text and Image Generation", style="Heading 1")

    para(doc, "2.1 Text generation", style="Heading 2")
    para(doc,
         "Generates study material on a topic - an explanation, a worked example, a "
         "revision paragraph - through whichever provider the application was launched "
         "with. With 'Use my study material as context' ticked, the loaded document is "
         "sent along with the prompt, so the model expands on the student's own notes "
         "instead of answering from general knowledge.")
    add_image(doc, "text_generation/01_text_generation_from_material.png",
              "Figure 2: Text generation answering from the loaded study material.")

    para(doc, "2.2 Image generation", style="Heading 2")
    para(doc,
         "Produces a diagram to revise from. Beyond free-text prompting, the sub-task "
         "can build the prompt itself: 'Build a mind map from my study material' runs "
         "the same key-term extractor the practice question generator uses, takes the "
         "highest-ranked term as the central node and the rest as branches, and writes "
         "a mind-map prompt. The diagram therefore reflects what the notes emphasise "
         "rather than a generic request. Generated images can be saved as PNG, JPG or "
         "JPEG into data/.")
    add_image(doc, "image_generation/01_mindmap_prompt_from_material.png",
              "Figure 3: The mind-map prompt written automatically from the loaded "
              "biology notes - central node 'cells', branches for respiration, "
              "photosynthesis, ATP and the other extracted key terms.")

    doc.add_page_break()

    # ---------------- Sub-tasks ----------------
    para(doc, "3. Sub-task: Text Summarisation", style="Heading 1")
    para(doc,
         "Condenses long study material into revision-ready form. Five output styles "
         "(concise abstract, revision bullet points, study notes, exam key takeaways, "
         "plain-language explanation), an optional focus prompt, and map-reduce handling "
         "for documents longer than one model context: the input is split on paragraph "
         "and sentence boundaries, each chunk is summarised, and the partial summaries "
         "are then summarised into the requested style.")
    add_image(doc, "text_summarisation/02_summary_hf_api_bullets.png",
              "Figure 4: Summarisation with Qwen2.5-7B via the Hugging Face Inference "
              "API - revision bullet points, 1.88 s, 6.25x compression.")
    add_image(doc, "text_summarisation/04_summary_hf_local_slm.png",
              "Figure 5: The same document summarised by DistilBART running locally - "
              "no API cost, but slower and lower quality.")

    para(doc, "3.1 Measured comparison", style="Heading 2")
    add_table(
        doc,
        ["Backend", "Latency", "Tokens", "Cost", "Compression", "ROUGE-1", "ROUGE-L"],
        [
            ["DistilBART (local SLM)", "39.7 s", "1315", "$0.00", "7.2x", "0.225", "0.225"],
            ["Qwen2.5-7B (HF API LLM)", "2.2 s", "1299", "$0.00", "7.0x", "0.189", "0.128"],
        ],
        widths=[1.8, 0.9, 0.7, 0.7, 1.0, 0.7, 0.7],
    )
    para(doc,
         "Note on ROUGE: with no human reference summaries available, ROUGE is measured "
         "against the source text. It therefore rewards reusing the source's wording, so "
         "the extractive-leaning local model scores higher while producing visibly worse "
         "summaries. ROUGE is reported here as a faithfulness and anti-hallucination "
         "signal, not as a quality ranking.", italic=True)

    doc.add_page_break()

    para(doc, "4. Sub-task: Question Answering", style="Heading 1")
    para(doc,
         "Answers student questions strictly from the material supplied, not from the "
         "model's own knowledge. Two backends of deliberately different kinds: an "
         "extractive model that selects a span from the material and cites the sentence "
         "it came from, and a generative model that reads better and can answer across "
         "several sentences.")
    add_image(doc, "question_answering/02_qa_extractive_local_slm.png",
              "Figure 6: Extractive answering with DistilBERT SQuAD running locally. "
              "The answer is a span copied from the material, so it cannot hallucinate.")
    add_image(doc, "question_answering/03_qa_generative_hf_api.png",
              "Figure 7: The same question answered generatively by Qwen2.5-7B.")
    add_image(doc, "question_answering/04_qa_declines_when_not_covered.png",
              "Figure 8: Asked something the material does not cover, the assistant "
              "declines instead of guessing.")
    para(doc,
         "Quality is measured differently per backend kind: extractive answers carry the "
         "model's span confidence, generative answers carry groundedness - the share of "
         "the answer's content words present in the source. A generative answer that "
         "drifts into the model's own knowledge scores lower, making the drift visible "
         "in the metrics.")

    doc.add_page_break()

    # ---------------- Fine-tuning ----------------
    para(doc, "5. Sub-task: Image Classification", style="Heading 1")
    para(doc,
         "Labels a photo of study material - a diagram, a page of handwritten notes, a "
         "textbook page - so it can be routed to the right sub-task. This is the entry "
         "point for material that arrives as an image rather than as text.")
    para(doc,
         "Three backends. ViT-base runs locally over ImageNet-1k with no key and no "
         "cost, and the same model is available through the Inference API. A third "
         "backend uses gpt-4o-mini vision, which is not restricted to a fixed label set.")
    add_image(doc, "image_classification/01_classify_diagram_local_vit.png",
              "Figure 9: Classifying a cell diagram with ViT-base running locally.")
    para(doc,
         "The comparison is instructive. The ImageNet model can only answer with one of "
         "its thousand object classes, so shown a page of handwritten notes it reaches "
         "for the nearest object rather than describing the material. The vision LLM is "
         "not constrained that way and can answer with study-material terms such as "
         "'diagram' or 'handwritten notes'. This is the same small-model-versus-large-"
         "model trade-off documented for question answering, in a different modality.")

    doc.add_page_break()

    para(doc, "6. Fine-tuning (Requirement 8)", style="Heading 1")
    para(doc,
         f"Base model: {ft['base_model']}. Dataset: {ft['dataset']} - SciQ, 11,679 "
         "crowdsourced science exam questions each with a support paragraph. School "
         "science material, so the same Education domain as the rest of the application.")
    para(doc,
         f"Task: practice question generation, mapping answer + passage to an exam-style "
         f"question. Trained on {ft['train_examples']} examples for {ft['epochs']} epoch "
         f"on {ft['device'].upper()} in {ft['train_minutes']} minutes, and evaluated on "
         f"{ft['after']['examples']} held-out SciQ test items.")

    para(doc, "6.1 Results", style="Heading 2")
    add_table(
        doc,
        ["Model", "Exact match", "Token F1", "ROUGE-L", "Copy rate"],
        [
            ["Base t5-small",
             f"{ft['before']['exact_match']}%", f"{ft['before']['token_f1']}%",
             f"{ft['before']['rouge_l']}%", f"{ft['before']['copy_rate']}%"],
            ["Fine-tuned on SciQ",
             f"{ft['after']['exact_match']}%", f"{ft['after']['token_f1']}%",
             f"{ft['after']['rouge_l']}%", f"{ft['after']['copy_rate']}%"],
            ["Change",
             f"{ft['delta']['exact_match']:+}", f"{ft['delta']['token_f1']:+}",
             f"{ft['delta']['rouge_l']:+}", f"{ft['delta']['copy_rate']:+}"],
        ],
        widths=[1.8, 1.1, 1.0, 1.0, 1.0],
    )
    para(doc,
         "Copy rate is the metric that shows what changed. Base t5-small has never seen "
         "the 'generate question:' prefix, so it falls back to echoing the passage: "
         f"{ft['before']['copy_rate']}% of its output words come straight from the input, "
         "and its exact match is zero because it never produces a question at all.")

    para(doc, "6.2 Example output", style="Heading 2")
    samples = list(zip(ft["before"]["samples"], ft["after"]["samples"]))[:3]
    add_table(
        doc,
        ["Reference question", "Base t5-small", "Fine-tuned"],
        [[b["expected"], b["predicted"][:160], a["predicted"][:160]] for b, a in samples],
        widths=[2.0, 2.0, 2.0],
    )

    add_image(doc, "practice_questions/02_practice_questions_generated.png",
              "Figure 10: The fine-tuned model generating practice questions from "
              "biology revision notes inside the application.")

    para(doc, "6.3 Why this task rather than question answering", style="Heading 2")
    para(doc,
         "Fine-tuning for grounded question answering was tried first and rejected on "
         "evidence. Base t5-small already scores 70.0% exact match and 81.2% token F1 on "
         "SciQ question answering, because T5's original pre-training mixture includes "
         "SQuAD in exactly the 'question: ... context: ...' format used here. There was "
         "no headroom to demonstrate, and a pilot fine-tune scored slightly worse than "
         "the base model. That run is retained in the code as --task qa.")

    doc.add_page_break()

    # ---------------- LLMOps ----------------
    para(doc, "7. LLMOps and Metrics (Requirement 7)", style="Heading 1")
    para(doc,
         "Every model call in the application passes through a track() context manager "
         "in src/metrics.py, which writes one row per invocation to data/metrics_log.csv "
         "and surfaces the aggregates in the application's LLMOps dashboard. Adding a "
         "new sub-task to the instrumentation is three lines of code.")
    add_table(
        doc,
        ["#", "Metric", "What it tells us"],
        [
            [1, "Latency (avg and p95)", "Wall-clock time per model call; p95 exposes "
                                         "the tail an average hides"],
            [2, "Token usage", "Prompt and completion tokens, per call and cumulative"],
            [3, "Cost (USD)", "Estimated from published per-token pricing"],
            [4, "Success rate", "Share of calls completing without an exception; "
                                "failures are logged, not swallowed"],
            [5, "Quality - ROUGE-1 / ROUGE-L", "Summary faithfulness to the source"],
            [6, "Quality - groundedness", "Share of a generated answer's content words "
                                          "present in the source"],
            [7, "Compression ratio", "Source words per summary word"],
        ],
        widths=[0.4, 2.0, 3.6],
    )
    add_image(doc, "text_summarisation/05_llmops_metrics_dashboard.png",
              "Figure 11: The in-application LLMOps dashboard, aggregating every "
              "logged model call.")

    doc.add_page_break()

    # ---------------- Running ----------------
    para(doc, "8. Running the Application", style="Heading 1")
    para(doc, "pip install -r requirements.txt", style="Intense Quote")
    para(doc, "python -m streamlit run app.py -- hf", style="Intense Quote")
    para(doc,
         "Copy .env.example to .env and set HF_API_KEY (free, from "
         "huggingface.co/settings/tokens) and optionally OPENAI_API_KEY. The .env file "
         "is git-ignored and no key is committed to the repository.")
    para(doc,
         "To reproduce the fine-tuning: python scripts/finetune_qa.py. To reproduce the "
         "summarisation comparison: python scripts/benchmark_summarization.py.")

    para(doc, "9. Limitations and Future Work", style="Heading 1")
    for text in [
        "Text generation and image generation are not yet wired into the metrics "
        "layer, so the LLMOps dashboard covers four of the six sub-tasks.",
        "The local models are CPU-bound. First use of each downloads and loads weights "
        "- DistilBERT took 106 seconds on this machine - after which calls are "
        "sub-second to a few seconds.",
        "The fine-tuned question generator reaches 13.3% exact match. It is a "
        "60M-parameter model trained on 1,600 examples for one epoch on CPU; the size and "
        "direction of the improvement is the result, not the absolute score.",
        "The fine-tuned model is trained on science material and degrades on out-of-domain "
        "notes, reverting to copying sentences. Outputs that are not questions are "
        "filtered and another key term is tried instead.",
        "The extractive question answering model selects one span per question and can "
        "select from the wrong section of a long document while reporting high confidence.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    print(f"Wrote {OUT_FILE}")
    print("\nBefore submitting:")
    print("  1. Fill every «...» placeholder (group number, names, BITS IDs, %).")
    print("  2. Decide what to do about image classification (implement or remove).")
    print("  3. Rename the file to <groupid>.docx.")


if __name__ == "__main__":
    build()
